import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def pdf_to_word(pdf_file, word_file):
    # 打开PDF文件
    doc = fitz.open(pdf_file)
    # 创建Word文档
    word_doc = Document()

    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        blocks = page.get_text("dict")["blocks"]

        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    paragraph = word_doc.add_paragraph()
                    for span in line["spans"]:
                        run = paragraph.add_run(span["text"])
                        run.font.size = Pt(span["size"])
                        if span["flags"] & 2:  # 检查粗体（bold）
                            run.bold = True
                        if span["flags"] & 1:  # 检查斜体（italic）
                            run.italic = True

                    # 处理对齐方式
                    if line["dir"][0] == 1:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif line["dir"][0] == -1:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    elif line["dir"][1] == 1:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 保存Word文档
    word_doc.save(word_file)
    print(f"PDF 文件已成功转换为 Word 文件：{word_file}")

# 使用示例：将 'input.pdf' 替换为你的PDF文件路径，'output.docx' 替换为你希望保存的Word文件路径
pdf_to_word('2024.pdf', 'output.docx')
